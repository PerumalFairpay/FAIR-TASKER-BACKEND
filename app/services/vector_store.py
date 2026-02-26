import os
import requests
import asyncio
from typing import List, Optional
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.core.config import QDRANT_URL, QDRANT_API_KEY, QDRANT_COLLECTION_NAME
import pypdf
from docx import Document as DocxDocument
import io
import logging

logger = logging.getLogger(__name__)

class VectorStoreService:
    def __init__(self):
        self.api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001",
            google_api_key=self.api_key
        )
        
        # Detect vector dimension dynamically to avoid hardcoding issues
        try:
            # We do a single test embedding to get the dimension
            test_vector = self.embeddings.embed_query("dimension_test")
            self.dimension = len(test_vector)
            logger.info(f"Detected embedding dimension: {self.dimension}")
        except Exception as e:
            logger.warning(f"Failed to detect dimension dynamically: {e}. Falling back to 3072.")
            self.dimension = 3072
        
        if QDRANT_URL and QDRANT_API_KEY:
            try:
                self.client = QdrantClient(
                    url=QDRANT_URL,
                    api_key=QDRANT_API_KEY,
                    prefer_grpc=True,
                )
                
                # Ensure collection exists
                collection_name = QDRANT_COLLECTION_NAME or "documents"
                allow_recreate = os.environ.get("QDRANT_ALLOW_RECREATE", "false").lower() == "true"
                
                try:
                    info = self.client.get_collection(collection_name=collection_name)
                    # Check for dimension mismatch
                    if info.config.params.vectors.size != self.dimension:
                        if allow_recreate:
                            logger.warning(f"Dimension mismatch (expected {self.dimension}, got {info.config.params.vectors.size}). Recreating collection.")
                            self.client.delete_collection(collection_name=collection_name)
                            raise ValueError("Recreate")
                        else:
                            logger.error(f"Dimension mismatch detected (expected {self.dimension}, got {info.config.params.vectors.size}). "
                                         "Recreation is disabled (QDRANT_ALLOW_RECREATE=false). Vector store will be unavailable.")
                            self.client = None
                            self.vector_store = None
                            return
                except Exception as e:
                    if str(e) == "Recreate" or "not found" in str(e).lower():
                        logger.info(f"Creating collection: {collection_name}")
                        from qdrant_client.http import models as rest
                        self.client.create_collection(
                            collection_name=collection_name,
                            vectors_config=rest.VectorParams(
                                size=self.dimension,
                                distance=rest.Distance.COSINE,
                            ),
                        )
                    else:
                        raise e
                    
                # Ensure payload indexes exist for filtering
                try:
                    from qdrant_client.http import models as rest
                    self.client.create_payload_index(
                        collection_name=collection_name,
                        field_name="metadata.document_id",
                        field_schema=rest.PayloadSchemaType.KEYWORD,
                    )
                    self.client.create_payload_index(
                        collection_name=collection_name,
                        field_name="metadata.category_id",
                        field_schema=rest.PayloadSchemaType.KEYWORD,
                    )
                    logger.info("Payload indices created successfully.")
                except Exception as e:
                    # Ignore if index already exists
                    if "already exists" not in str(e).lower():
                        logger.warning(f"Note: Payload index creation skipped or failed: {e}")

                self.vector_store = QdrantVectorStore(
                    client=self.client,
                    collection_name=collection_name,
                    embedding=self.embeddings,
                )
            except Exception as e:
                logger.error(f"Failed to initialize Qdrant Cloud: {str(e)}")
                self.client = None
                self.vector_store = None
        else:
            self.client = None
            self.vector_store = None
            logger.warning("Qdrant Cloud configuration is missing. Vector store will not be available.")

    def _extract_text(self, file_content: bytes, filename: str, file_type: str = None) -> str:
        """Helper to extract text from file content (CPU bound)."""
        text = ""
        if file_type == "application/pdf" or filename.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(file_content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            try:
                text = file_content.decode("utf-8")
            except:
                return ""
        return text

    async def index_document(self, file_url: str, metadata: dict, file_type: str = None):
        """Extract text from file and index it in Qdrant (Non-blocking)."""
        if not self.vector_store:
            logger.warning("Vector store not initialized, skipping indexing.")
            return

        # Metadata Validation
        document_id = metadata.get("document_id")
        if not document_id:
            logger.error(f"Missing document_id in metadata for {file_url}. Skipping indexing.")
            return

        try:
            logger.info(f"Starting indexing for document: {metadata.get('name')} (ID: {document_id})")
            
            # Idempotency: Delete existing vectors for this document_id first
            # This prevents duplicate chunks if the document is re-uploaded or updated
            await self.delete_document(document_id)
            
            # 1. Get file content directly from storage
            from app.helper.file_handler import file_handler
            file_id = file_url.split("/")[-1]
            
            # Offload blocking file reading to thread
            file_data = await asyncio.to_thread(file_handler.get_file, file_id)
            
            file_content = None
            if isinstance(file_data, dict) and "Body" in file_data:
                # S3 Storage (blocking read)
                file_content = await asyncio.to_thread(file_data["Body"].read)
                logger.info("Fetched file content from S3.")
            else:
                # Check for local file path
                local_info = await asyncio.to_thread(file_handler.get_file_info, file_id)
                if local_info and os.path.exists(local_info):
                    # Blocking file read
                    def read_local():
                        with open(local_info, "rb") as f:
                            return f.read()
                    file_content = await asyncio.to_thread(read_local)
                    logger.info(f"Read file content from local path: {local_info}")
                else:
                    # Final fallback to requests
                    logger.info(f"Attempting fallback download via HTTP: {file_url}")
                    response = await asyncio.to_thread(requests.get, file_url, timeout=10)
                    if response.status_code == 200:
                        file_content = response.content
                    else:
                        logger.error(f"Failed to fetch file. Status: {response.status_code}")
                        return

            if not file_content:
                logger.error("Could not retrieve file content for indexing.")
                return

            # 2. Extract text (CPU bound)
            filename = metadata.get("name", "").lower() or file_url.lower()
            text = await asyncio.to_thread(self._extract_text, file_content, filename, file_type)

            if not text or not text.strip():
                logger.warning(f"No text extracted from document: {metadata.get('name')}")
                return

            logger.info(f"Extracted {len(text)} characters. Chunking...")

            # 3. Chunk text (CPU bound)
            def chunk_text():
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000, 
                    chunk_overlap=200
                )
                return text_splitter.split_text(text)
            
            chunks = await asyncio.to_thread(chunk_text)
            logger.info(f"Created {len(chunks)} chunks.")

            # 4. Create documents with metadata
            docs = [
                Document(page_content=chunk, metadata={**metadata, "chunk_index": i})
                for i, chunk in enumerate(chunks)
            ]

            # 5. Add to Qdrant (Blocking I/O)
            await asyncio.to_thread(self.vector_store.add_documents, docs)
            logger.info(f"Successfully indexed document: {metadata.get('name')} in Qdrant.")

        except Exception as e:
            logger.error(f"Error indexing document {metadata.get('name')}: {str(e)}", exc_info=True)

    async def search_documents(self, query: str, filter_dict: dict = None, limit: int = 5) -> List[dict]:
        """Search for relevant document snippets (Non-blocking)."""
        if not self.vector_store:
            return []

        try:
            # Prepare Qdrant filter if filter_dict is provided
            qdrant_filter = None
            if filter_dict:
                from qdrant_client.http import models as rest
                must_conditions = []
                for key, value in filter_dict.items():
                    # We assume these are keyword matches for metadata
                    # Note: LangChain stores metadata with 'metadata.' prefix in Qdrant payload
                    # But the LangChain Qdrant integration usually handles the mapping if passed to its filter param
                    must_conditions.append(
                        rest.FieldCondition(
                            key=f"metadata.{key}",
                            match=rest.MatchValue(value=value),
                        )
                    )
                if must_conditions:
                    qdrant_filter = rest.Filter(must=must_conditions)

            # Offload blocking similarity search to thread
            # The 'filter' parameter is supported by LangChain's Qdrant similarity_search
            results = await asyncio.to_thread(
                self.vector_store.similarity_search, 
                query, 
                k=limit,
                filter=qdrant_filter
            )
            
            return [
                {
                    "content": res.page_content,
                    "metadata": res.metadata
                }
                for res in results
            ]
        except Exception as e:
            logger.error(f"Search failed: {str(e)}")
            return []

    async def delete_document(self, document_id: str):
        """Delete all vectors associated with a document ID (Non-blocking)."""
        if not self.client or not document_id:
            return

        try:
            from qdrant_client.http import models as rest
            
            def run_delete():
                self.client.delete(
                    collection_name=QDRANT_COLLECTION_NAME,
                    points_selector=rest.Filter(
                        must=[
                            rest.FieldCondition(
                                key="metadata.document_id",
                                match=rest.MatchValue(value=document_id),
                            ),
                        ]
                    ),
                )
            
            await asyncio.to_thread(run_delete)
            logger.info(f"Deleted vectors for document_id: {document_id}")
        except Exception as e:
            logger.error(f"Failed to delete vectors: {str(e)}")

vector_store_service = VectorStoreService()
